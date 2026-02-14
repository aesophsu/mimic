import os
import glob
from docx import Document
from docx.shared import Inches


PROJECT_ROOT = os.path.normpath(os.path.join(os.path.dirname(__file__), "..", ".."))
DOCS_DIR = os.path.join(PROJECT_ROOT, "docs")
SUPP_DOCS_DIR = os.path.join(DOCS_DIR, "supplementary")
SUPP_FIG_DIR = os.path.join(DOCS_DIR, "figures", "supplementary")


def _ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_separator(doc: Document) -> None:
    # 用一行短横线模拟「⸻」分隔
    doc.add_paragraph("—" * 10)


def _pick_first_png(pattern: str) -> str | None:
    paths = sorted(glob.glob(pattern))
    if not paths:
        return None
    # 优先 png
    pngs = [p for p in paths if p.lower().endswith(".png")]
    return pngs[0] if pngs else paths[0]


def add_figure(doc: Document, img_path: str | None, width_in: float = 6.0) -> None:
    if not img_path or not os.path.exists(img_path):
        return
    doc.add_picture(img_path, width=Inches(width_in))


def build_s1_docx() -> None:
    """S1 — LASSO Feature Selection"""
    _ensure_dir(SUPP_DOCS_DIR)
    doc = Document()

    add_heading(doc, "Supplementary Material S1 — LASSO Feature Selection", level=1)
    add_paragraph(doc, "Supplementary Figure S1. LASSO coefficient paths and selected predictors")
    add_paragraph(doc, "Panel S1A. LASSO coefficient paths")
    add_paragraph(
        doc,
        "This panel shows the coefficient shrinkage trajectories across a range of penalty "
        "values (λ). As λ increases, coefficients progressively shrink toward zero, with the "
        "final set of predictors determined at the cross-validated optimal λ for each outcome "
        "(POF, 28-day mortality, and the composite endpoint).",
    )

    sf1_dir = os.path.join(SUPP_FIG_DIR, "SF1_lasso")
    # Panel S1A: coefficient paths for POF, mortality, composite
    for outcome in ("pof", "mortality", "composite"):
        fig = _pick_first_png(os.path.join(sf1_dir, f"SF1a_diag_{outcome}.*"))
        add_figure(doc, fig)
    add_separator(doc)

    add_paragraph(doc, "Panel S1B. Selected predictors and standardized coefficients")
    add_paragraph(
        doc,
        "This panel displays predictors retained at the optimal λ and their standardized "
        "coefficients. Positive coefficients indicate higher predicted risk, and negative "
        "values indicate protective associations. Selected predictors primarily reflect "
        "renal function, oxygenation, perfusion, and acid–base status.",
    )
    # Panel S1B: importance plots
    for outcome in ("pof", "mortality", "composite"):
        fig = _pick_first_png(os.path.join(sf1_dir, f"SF1b_importance_{outcome}.*"))
        add_figure(doc, fig)

    out_path = os.path.join(SUPP_DOCS_DIR, "S1_LASSO_Feature_Selection.docx")
    doc.save(out_path)


def build_s2_docx() -> None:
    """S2 — Internal Validation"""
    _ensure_dir(SUPP_DOCS_DIR)
    doc = Document()

    add_heading(doc, "Supplementary Material S2 — Internal Validation", level=1)
    add_paragraph(doc, "Supplementary Figure S2. ROC and calibration curves in the internal validation cohort")

    add_paragraph(doc, "Panel S2A. ROC curves")
    add_paragraph(
        doc,
        "ROC curves for POF, 28-day mortality, and the composite endpoint in the MIMIC-IV "
        "internal validation cohort are shown. All models demonstrated good discriminative "
        "performance, with AUC values typically ranging from 0.82 to 0.88.",
    )

    sf2_dir = os.path.join(SUPP_FIG_DIR, "SF2_internal_ROC")
    for outcome in ("pof", "mortality", "composite"):
        fig = _pick_first_png(os.path.join(sf2_dir, f"SF2a_ROC_{outcome}.*"))
        add_figure(doc, fig)
    add_separator(doc)

    add_paragraph(doc, "Panel S2B. Calibration curves")
    add_paragraph(
        doc,
        "Calibration plots compare predicted and observed risks. Curves close to the reference "
        "line indicate good calibration; mild underestimation was observed in the high-risk range.",
    )
    for outcome in ("pof", "mortality", "composite"):
        fig = _pick_first_png(os.path.join(sf2_dir, f"SF2b_Calibration_{outcome}.*"))
        add_figure(doc, fig)

    out_path = os.path.join(SUPP_DOCS_DIR, "S2_Internal_Validation.docx")
    doc.save(out_path)


def build_s3_docx() -> None:
    """S3 — Diagnostic Performance & Extended Interpretability"""
    _ensure_dir(SUPP_DOCS_DIR)
    doc = Document()

    add_heading(doc, "Supplementary Material S3 — Diagnostic Performance and Extended Interpretability", level=1)
    add_paragraph(doc, "Supplementary Figure S3. Diagnostic metrics, extended feature importance, and forest plots")

    sf3_dir = os.path.join(SUPP_FIG_DIR, "SF3_cutoff")

    # Panel S3A
    add_paragraph(doc, "Panel S3A. Diagnostic statistics")
    add_paragraph(
        doc,
        "This panel summarizes sensitivity, specificity, and predictive values at the optimal "
        "Youden-index threshold for each model, illustrating diagnostic performance across "
        "clinically relevant cut-points.",
    )
    for outcome in ("pof", "mortality", "composite"):
        fig = _pick_first_png(os.path.join(sf3_dir, f"SF3a_Diagnostic_{outcome}.*"))
        add_figure(doc, fig)
    add_separator(doc)

    # Panel S3B
    add_paragraph(doc, "Panel S3B. Extended feature importance")
    add_paragraph(
        doc,
        "Feature importance derived from tree-based models (random forest and XGBoost) is "
        "presented. The most influential features align with LASSO-selected variables and "
        "reflect key physiologic domains.",
    )
    imp_fig = _pick_first_png(os.path.join(sf3_dir, "SF3b_feature_importance.*"))
    add_figure(doc, imp_fig)
    add_separator(doc)

    # Panel S3C
    add_paragraph(doc, "Panel S3C. Forest plot")
    add_paragraph(
        doc,
        "Multivariable logistic regression–based odds ratios with 95% confidence intervals are "
        "shown for major predictors of POF. Renal dysfunction, impaired oxygenation, and "
        "acid–base abnormalities were consistently associated with higher risk.",
    )
    forest_fig = _pick_first_png(os.path.join(sf3_dir, "SF3c_forest_plot.*"))
    add_figure(doc, forest_fig)

    out_path = os.path.join(SUPP_DOCS_DIR, "S3_Diagnostic_Performance_and_Interpretability.docx")
    doc.save(out_path)


def build_s4_docx() -> None:
    """S4 — Dataset Shift Analyses"""
    _ensure_dir(SUPP_DOCS_DIR)
    doc = Document()

    add_heading(doc, "Supplementary Material S4 — Dataset Shift Analyses (MIMIC-IV vs. eICU)", level=1)
    add_paragraph(doc, "Supplementary Figure S4. Distributional comparison between development and validation cohorts")

    sf4_dir = os.path.join(SUPP_FIG_DIR, "SF4_comparison")

    # Panel S4A
    add_paragraph(doc, "Panel S4A. Cross-cohort distribution of key variables")
    add_paragraph(
        doc,
        "This panel compares distributions of major baseline variables between MIMIC-IV and "
        "eICU cohorts to illustrate the degree of population- and practice-related variation "
        "relevant to external validation.",
    )
    fig_a = _pick_first_png(os.path.join(sf4_dir, "SF4a_Table4_visualization.*"))
    add_figure(doc, fig_a)
    add_separator(doc)

    # Panel S4B
    add_paragraph(doc, "Panel S4B. Variable-wise drift analyses")
    add_paragraph(
        doc,
        "Density plots, cumulative distribution functions, and distance metrics depict "
        "variable-specific distributional drift across cohorts. Most drift arises from "
        "differences in practice patterns, such as ventilation use and laboratory sampling "
        "frequency. The following panels (S4B1–S4B[N]) show variable-wise drift analyses for "
        "each key predictor.",
    )
    drift_dir = os.path.join(sf4_dir, "SF4b_drift")
    drift_figs = sorted(glob.glob(os.path.join(drift_dir, "SF4b_*.png")))
    for fig in drift_figs:
        add_figure(doc, fig, width_in=5.5)

    out_path = os.path.join(SUPP_DOCS_DIR, "S4_Dataset_Shift_Analyses.docx")
    doc.save(out_path)


def build_s5_docx() -> None:
    """S5 — SHAP-Based Explainability"""
    _ensure_dir(SUPP_DOCS_DIR)
    doc = Document()

    add_heading(doc, "Supplementary Material S5 — SHAP-Based Explainability", level=1)
    add_paragraph(doc, "Supplementary Figure S5. Individual-level SHAP interpretation")

    sf5_dir = os.path.join(SUPP_FIG_DIR, "SF5_interpretation")

    # Panel S5A
    add_paragraph(doc, "Panel S5A. SHAP force plots")
    add_paragraph(
        doc,
        "Force plots illustrate individual prediction decompositions. Red elements denote "
        "features pushing risk upward, whereas blue elements indicate risk-lowering effects.",
    )
    for outcome in ("pof", "mortality", "composite"):
        fig = _pick_first_png(os.path.join(sf5_dir, f"SF5a_Force_{outcome}.*"))
        add_figure(doc, fig)
    add_separator(doc)

    # Panel S5B
    add_paragraph(doc, "Panel S5B. SHAP dependence plots")
    add_paragraph(
        doc,
        "Dependence plots show how changes in key predictors (e.g., creatinine, BUN, P/F "
        "ratio, pH, lactate, PTT) influence model output. These patterns highlight the "
        "central role of renal function, oxygenation, perfusion, and acid–base status in "
        "early deterioration.",
    )
    dep_dir = os.path.join(sf5_dir, "SF5b_Dep")
    dep_figs = sorted(glob.glob(os.path.join(dep_dir, "SF5b_Dep_*.png")))
    for fig in dep_figs:
        add_figure(doc, fig, width_in=5.5)

    out_path = os.path.join(SUPP_DOCS_DIR, "S5_SHAP_Based_Explainability.docx")
    doc.save(out_path)


def main() -> None:
    build_s1_docx()
    build_s2_docx()
    build_s3_docx()
    build_s4_docx()
    build_s5_docx()


if __name__ == "__main__":
    main()

